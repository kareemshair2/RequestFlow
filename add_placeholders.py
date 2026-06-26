import os
from docx import Document
from docx.oxml.ns import qn

templates_dir = r'C:\Users\nour\OneDrive\Desktop\templates'
output_dir = r'C:\Users\nour\OneDrive\Desktop\templates\modified'
os.makedirs(output_dir, exist_ok=True)

def smart_replace(para, old_text, new_text):
    full = para.text
    if old_text not in full:
        return False
    replaced = full.replace(old_text, new_text)
    if para.runs:
        first = para.runs[0]
        for r in para.runs[1:]:
            r.text = ''
        first.text = replaced
    else:
        para.add_run(replaced)
    return True

def process_paragraphs(doc, replacements):
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if smart_replace(para, old, new):
                break

def process_tables(doc, replacements):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if smart_replace(para, old, new):
                            break

def clear_table_rows(table, keep_header=True, keep_first_data=False):
    """Remove all data rows from table, keeping header and optionally one data row"""
    rows_to_remove = []
    start_idx = 1  # Skip header row
    
    if keep_first_data and len(table.rows) > 1:
        # Clear content of first data row
        for cell in table.rows[1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = ''
        start_idx = 2
    
    for idx in range(len(table.rows) - 1, start_idx - 1, -1):
        rows_to_remove.append(table.rows[idx])
    
    tbl = table._tbl
    for row in rows_to_remove:
        tbl.remove(row._tr)

# ============ Purchase Template ============
print('=== تمبلت الشراء ===')
doc = Document(os.path.join(templates_dir, 'تمبلت الشراء.docx'))
process_paragraphs(doc, {
    'تاريخ الطلب: ': 'تاريخ الطلب: {{requestDate}}  ',
    'الإدارة الطالبة:                                                                                                               المشروع: ':
        'الإدارة الطالبة: {{requestingDept}}        المشروع: {{project}}   ',
    'موضوع الطلب:                                                                                                              الوظيفة: ':
        'موضوع الطلب: {{subject}}        الوظيفة: {{jobTitle}}   ',
    'مقدم الطلب: ': 'مقدم الطلب: {{requester}}  ',
    'تفاصيل الطلب:-\nتوقيع الجهة الطالبة:  ':
        'تفاصيل الطلب:- {{details}}    توقيع الجهة الطالبة: {{requesterSignature}} ',
    'مراجعة المخازن:': 'مراجعة المخازن: {{warehouseReview}}  ',
    'الإدارة المالية لسماح البند:  ':
        'الإدارة المالية لسماح البند: {{financialApproval}}  ',
    'المدير التنفيذي للموافقة:': 'المدير التنفيذي للموافقة: {{executiveApproval}}  ',
})
doc.save(os.path.join(output_dir, 'تمبلت الشراء.docx'))
print('  Done')

# ============ Disbursement Template ============
print('=== تمبلت طلبات الصرف ===')
doc = Document(os.path.join(templates_dir, 'تمبلت طلبات الصرف.docx'))
process_paragraphs(doc, {
    'اقر انا / …………………………………..':
        'اقر انا / {{recipientName}}  ',
    'رقم قومي /': 'رقم قومي / {{nationalId}}  ',
    'انني قد استلمت من جمعية صناع الحياة المنوفية منحية عينية عبارة عن كتب مدرسية للصف الثالث الثانوي':
        'انني قد استلمت من جمعية صناع الحياة المنوفية منحية عينية عبارة عن: {{acknowledgment}}',
    'الاسم /': 'الاسم / {{signatureName}}',
    'التوقيع /': 'التوقيع / {{signature}}',
})
doc.save(os.path.join(output_dir, 'تمبلت طلبات الصرف.docx'))
print('  Done')

# ============ Cases Template ============
print('=== كشف الحالات ===')
doc = Document(os.path.join(templates_dir, 'كشف الحالات.docx'))
# Replace title in table
process_tables(doc, {
    'كشف توزيع توزيع طقم لحاف وبطانية - جمعية صناع الحياة المنوفية': '{{title}}',
})

# Keep title row + header row + 1 empty data row
for table in doc.tables:
    while len(table.rows) > 3:
        tbl = table._tbl
        tr = table.rows[-1]._tr
        tbl.remove(tr)
    # Clear the single remaining data row (row 2)
    if len(table.rows) > 2:
        for cell in table.rows[2].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = ''
    # Title row has merged cells - don't clear them (they share content)

doc.save(os.path.join(output_dir, 'كشف الحالات.docx'))
print(f'  Done (cleared to header + 1 row)')

print('\nAll templates updated! Upload these 3 files to Drive folder "تمبلت الطلبات"')
